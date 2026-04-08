"""Tests for structured unit outline extraction."""

from pathlib import Path

import docx

from botstash.extractors.unit_outline import (
    _extract_assessments,
    _extract_learning_outcomes,
    _extract_metadata,
    _extract_weekly_schedule,
    _guess_bloom,
    _guess_category,
    _parse_outline,
    extract_unit_outline,
)

# -- Bloom's taxonomy guesser ------------------------------------------------


def test_guess_bloom_create() -> None:
    assert _guess_bloom("Design a database schema") == "create"


def test_guess_bloom_evaluate() -> None:
    assert _guess_bloom("Evaluate the effectiveness of AI") == "evaluate"


def test_guess_bloom_apply() -> None:
    assert _guess_bloom("Apply knowledge to solve problems") == "apply"


def test_guess_bloom_default() -> None:
    assert _guess_bloom("Some generic text") == "understand"


# -- Assessment category guesser --------------------------------------------


def test_category_exam() -> None:
    assert _guess_category("Final Exam") == "exam"


def test_category_project() -> None:
    assert _guess_category("Group Project Submission") == "project"


def test_category_report() -> None:
    assert _guess_category("Technical Report") == "report"


def test_category_default() -> None:
    assert _guess_category("Weekly Homework") == "assignment"


# -- Metadata extraction ----------------------------------------------------


def test_extract_metadata_full() -> None:
    text = (
        "ISYS2001 (V.1) Introduction to Information Systems\n"
        "Credit Points: 25\n"
        "Semester 1, 2026\n"
        "Pre-requisite units: COMP1001\n"
        "Mode of study: Internal\n"
    )
    meta = _extract_metadata(text)
    assert meta["unit_code"] == "ISYS2001"
    assert meta["unit_title"] == "Introduction to Information Systems"
    assert meta["credit_points"] == 25
    assert meta["year"] == 2026
    assert meta["semester"] == "Semester 1"
    assert meta["prerequisites"] == "COMP1001"
    assert meta["delivery_mode"] == "Internal"


def test_extract_metadata_minimal() -> None:
    text = "Some random document with no structured metadata."
    meta = _extract_metadata(text)
    assert meta["unit_code"] is None


# -- Learning outcomes -------------------------------------------------------


def test_extract_learning_outcomes() -> None:
    text = (
        "Some preamble text\n"
        "Learning Outcomes\n"
        "On successful completion of this unit student can:\n"
        "1\n"
        "Design a relational database schema\n"
        "2\n"
        "Evaluate the effectiveness of information systems\n"
        "Assessment Schedule\n"
        "Next section content\n"
    )
    outcomes = _extract_learning_outcomes(text)
    assert len(outcomes) == 2
    assert outcomes[0].code == "ULO1"
    assert "database" in outcomes[0].description.lower()
    assert outcomes[0].bloom_level == "create"
    assert outcomes[1].bloom_level == "evaluate"


def test_extract_learning_outcomes_numbered_dot() -> None:
    """Fallback: '1. Description' format."""
    text = (
        "Learning Outcomes\n"
        "1. Explain the role of IS in organisations\n"
        "2. Apply systems thinking to business problems\n"
        "Assessment Schedule\n"
    )
    outcomes = _extract_learning_outcomes(text)
    assert len(outcomes) == 2
    assert outcomes[0].bloom_level == "understand"
    assert outcomes[1].bloom_level == "apply"


# -- Assessments -------------------------------------------------------------


def test_extract_assessments_structured() -> None:
    text = (
        "Assessment Schedule\n"
        "1\n"
        "Database Design Report\n"
        "30 %\n"
        "Week:5\n"
        "2\n"
        "Final Exam\n"
        "40 %\n"
        "Week:13\n"
        "Learning Outcomes\n"
    )
    assessments = _extract_assessments(text)
    assert len(assessments) == 2
    assert assessments[0].title == "Database Design Report"
    assert assessments[0].weight == 30.0
    assert assessments[0].due_week == 5
    assert assessments[0].category == "report"
    assert assessments[1].category == "exam"


def test_extract_assessments_fallback() -> None:
    """Fallback: 'Title NN%' format."""
    text = (
        "Assessment Summary\n"
        "Group Project 40%\n"
        "Final Exam 60%\n"
        "Weekly Schedule\n"
    )
    assessments = _extract_assessments(text)
    assert len(assessments) == 2
    assert assessments[0].weight == 40.0
    assert assessments[1].weight == 60.0


# -- Weekly schedule ---------------------------------------------------------


def test_extract_weekly_schedule() -> None:
    text = (
        "Weekly Schedule\n"
        "Week 1: Introduction to IS\n"
        "Week 2: Data and Information\n"
        "Week 3: Database Fundamentals\n"
        "Assessment Schedule\n"
    )
    weeks = _extract_weekly_schedule(text)
    assert len(weeks) == 3
    assert weeks[0].week_number == 1
    assert weeks[0].topic == "Introduction to IS"
    assert weeks[2].topic == "Database Fundamentals"


# -- Full parsing + formatting -----------------------------------------------


def test_parse_outline_structured() -> None:
    text = (
        "ISYS2001 (V.1) Introduction to IS\n"
        "Credit Points: 25\n"
        "Semester 1, 2026\n\n"
        "Learning Outcomes\n"
        "1. Explain the role of IS\n"
        "2. Design a database\n\n"
        "Assessment Schedule\n"
        "Assignment 1 30%\n"
        "Final Exam 70%\n\n"
        "Weekly Schedule\n"
        "Week 1: Introduction\n"
        "Week 2: Data\n"
    )
    data = _parse_outline(text)
    assert data.unit_code == "ISYS2001"
    assert len(data.learning_outcomes) == 2
    assert len(data.assessments) == 2
    assert len(data.weekly_schedule) == 2


def test_extract_unit_outline_docx(tmp_path: Path) -> None:
    """Full end-to-end: DOCX with structured content → markdown."""
    doc = docx.Document()
    doc.add_heading("COMP1001 Intro to Programming", level=1)
    doc.add_paragraph("Credit Points: 25")
    doc.add_paragraph("Semester 2, 2026")
    doc.add_heading("Learning Outcomes", level=2)
    doc.add_paragraph("1. Design algorithms for simple problems")
    doc.add_paragraph("2. Implement solutions in Python")
    doc.add_heading("Assessment Schedule", level=2)
    doc.add_paragraph("Programming Project 50%")
    doc.add_paragraph("Final Exam 50%")
    doc.add_heading("Weekly Schedule", level=2)
    doc.add_paragraph("Week 1: Introduction to Python")
    doc.add_paragraph("Week 2: Variables and Data Types")

    doc_path = tmp_path / "outline.docx"
    doc.save(str(doc_path))

    result = extract_unit_outline(str(doc_path))

    # Should have structured sections
    assert "# Unit Outline:" in result
    assert "## Assessments" in result
    assert "## Learning Outcomes" in result
    assert "## Weekly Schedule" in result
    assert "50%" in result
    assert "Week 1" in result
    assert "Design" in result


def test_extract_unit_outline_fallback(tmp_path: Path) -> None:
    """When no structure is detected, includes full text."""
    doc = docx.Document()
    doc.add_paragraph("Just some random text with no structured content.")
    doc.add_paragraph("More text that doesn't match any patterns.")

    doc_path = tmp_path / "outline.docx"
    doc.save(str(doc_path))

    result = extract_unit_outline(str(doc_path))
    assert "## Full Text" in result
    assert "random text" in result
